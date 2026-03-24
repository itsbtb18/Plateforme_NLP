import logging
import os
from collections.abc import Sequence
from typing import cast

from accounts.blocking import blocked_user_ids_for, exclude_hidden_users
from accounts.views import LoginAndVerifiedRequiredMixin
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import LoginRequiredMixin, UserPassesTestMixin
from django.db import models
from django.db.models import Q
from django.http import Http404, JsonResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import reverse, reverse_lazy
from django.utils.translation import gettext_lazy as _
from django.views.decorators.http import require_GET
from django.views.generic import DetailView, ListView
from django.views.generic.edit import DeleteView, FormView

from .forms import ResourceForm
from .models import (
    Article,
    Corpus,
    Course,
    Document,
    Memoir,
    NLPTool,
    ResourceBase,
    Thesis,
)

logger = logging.getLogger(__name__)

ResourceVariant = Document | NLPTool | Course | Corpus


class ResourceListView(LoginAndVerifiedRequiredMixin, ListView):
    template_name = "resources/list.html"
    context_object_name = "resources"
    paginate_by = 10

    def get_queryset(self) -> list[ResourceVariant]:
        search_query = self.request.GET.get("q", "")
        resource_type = self.request.GET.get("type", "")
        field_filter = self.request.GET.get("field", "")
        language_filter = self.request.GET.get("language", "")

        # STRICT: Public resources list only shows approved content
        # Pending resources are only visible in the admin pending review panel
        approval_filter = {"approval_status": "approved"}

        querysets = []

        if resource_type in ["", "article", "thesis", "memoir"]:
            docs = exclude_hidden_users(
                Document.objects.filter(**approval_filter),
                self.request.user,
                ("author",),
            )
            if language_filter:
                docs = docs.filter(language=language_filter)
            if resource_type in ["article", "thesis", "memoir"]:
                docs = docs.filter(document_type=resource_type)
            if search_query:
                docs = docs.filter(
                    Q(title__icontains=search_query)
                    | Q(description__icontains=search_query)
                    | Q(title_ar__icontains=search_query)
                    | Q(title_en__icontains=search_query)
                )
            querysets.append(docs)

        if resource_type in ["", "tool"]:
            tools = exclude_hidden_users(
                NLPTool.objects.filter(**approval_filter),
                self.request.user,
                ("author",),
            )
            if language_filter:
                tools = tools.filter(supported_languages__contains=language_filter)
            if search_query:
                tools = tools.filter(
                    Q(title__icontains=search_query)
                    | Q(description__icontains=search_query)
                    | Q(title_ar__icontains=search_query)
                    | Q(title_en__icontains=search_query)
                )
            querysets.append(tools)

        if resource_type in ["", "course"]:
            courses = exclude_hidden_users(
                Course.objects.filter(**approval_filter),
                self.request.user,
                ("author", "teacher"),
            )
            if language_filter:
                courses = courses.filter(language=language_filter)
            if field_filter:
                courses = courses.filter(field=field_filter)
            if search_query:
                courses = courses.filter(
                    Q(title__icontains=search_query)
                    | Q(description__icontains=search_query)
                    | Q(title_ar__icontains=search_query)
                    | Q(title_en__icontains=search_query)
                )
            querysets.append(courses)

        if resource_type in ["", "corpus"]:
            corpora = exclude_hidden_users(
                Corpus.objects.filter(**approval_filter), self.request.user, ("author",)
            )
            if language_filter:
                corpora = corpora.filter(language=language_filter)
            if field_filter:
                corpora = corpora.filter(field=field_filter)
            if search_query:
                corpora = corpora.filter(
                    Q(title__icontains=search_query)
                    | Q(description__icontains=search_query)
                    | Q(title_ar__icontains=search_query)
                    | Q(title_en__icontains=search_query)
                )
            querysets.append(corpora)

        combined: list[ResourceVariant] = []
        for qs in querysets:
            for obj in qs:
                obj.resource_type = self.get_resource_type(obj)
                combined.append(obj)

        # Handle sorting
        sort_by = self.request.GET.get("sort", "newest")
        if sort_by == "oldest":
            return sorted(combined, key=lambda x: x.creation_date, reverse=False)
        elif sort_by == "popular":
            return sorted(
                combined, key=lambda x: getattr(x, "views_count", 0), reverse=True
            )
        else:  # default: newest
            return sorted(combined, key=lambda x: x.creation_date, reverse=True)

    def get_resource_type(self, obj):
        if isinstance(obj, Document):
            return obj.document_type
        elif isinstance(obj, NLPTool):
            return "tool"
        elif isinstance(obj, Course):
            return "course"
        elif isinstance(obj, Corpus):
            return "corpus"
        return "unknown"

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["current_query"] = self.request.GET.urlencode()
        resources = cast(Sequence[ResourceVariant], self.object_list)
        context["total_count"] = len(resources)
        from .models import FieldChoices

        context["field_choices"] = FieldChoices.choices
        context["current_field"] = self.request.GET.get("field", "")
        context["current_language"] = self.request.GET.get("language", "")
        context["current_type"] = self.request.GET.get("type", "")
        context["current_sort"] = self.request.GET.get("sort", "newest")
        context["page"] = "resources"
        return context


@login_required
def resource_ajax_search(request):
    """AJAX endpoint for live search with JSON response."""
    search_query = request.GET.get("q", "").strip()
    resource_type = request.GET.get("type", "")
    sort_by = request.GET.get("sort", "newest")

    # STRICT: Public search only shows approved content
    approval_filter = {"approval_status": "approved"}
    querysets = []

    def get_resource_type_label(obj):
        if isinstance(obj, Document):
            return obj.document_type
        elif isinstance(obj, NLPTool):
            return "tool"
        elif isinstance(obj, Course):
            return "course"
        elif isinstance(obj, Corpus):
            return "corpus"
        return "unknown"

    if resource_type in ["", "article", "thesis", "memoir"]:
        docs = exclude_hidden_users(
            Document.objects.filter(**approval_filter), request.user, ("author",)
        )
        if resource_type in ["article", "thesis", "memoir"]:
            docs = docs.filter(document_type=resource_type)
        if search_query:
            docs = docs.filter(
                Q(title__icontains=search_query)
                | Q(description__icontains=search_query)
                | Q(title_ar__icontains=search_query)
                | Q(title_en__icontains=search_query)
                | Q(keywords__icontains=search_query)
            )
        querysets.append(docs)

    if resource_type in ["", "tool"]:
        tools = exclude_hidden_users(
            NLPTool.objects.filter(**approval_filter), request.user, ("author",)
        )
        if search_query:
            tools = tools.filter(
                Q(title__icontains=search_query)
                | Q(description__icontains=search_query)
                | Q(title_ar__icontains=search_query)
                | Q(title_en__icontains=search_query)
                | Q(keywords__icontains=search_query)
            )
        querysets.append(tools)

    if resource_type in ["", "course"]:
        courses = exclude_hidden_users(
            Course.objects.filter(**approval_filter),
            request.user,
            ("author", "teacher"),
        )
        if search_query:
            courses = courses.filter(
                Q(title__icontains=search_query)
                | Q(description__icontains=search_query)
                | Q(title_ar__icontains=search_query)
                | Q(title_en__icontains=search_query)
                | Q(keywords__icontains=search_query)
            )
        querysets.append(courses)

    if resource_type in ["", "corpus"]:
        corpora = exclude_hidden_users(
            Corpus.objects.filter(**approval_filter), request.user, ("author",)
        )
        if search_query:
            corpora = corpora.filter(
                Q(title__icontains=search_query)
                | Q(description__icontains=search_query)
                | Q(title_ar__icontains=search_query)
                | Q(title_en__icontains=search_query)
                | Q(keywords__icontains=search_query)
            )
        querysets.append(corpora)

    combined = []
    for qs in querysets:
        for obj in qs:
            rtype = get_resource_type_label(obj)
            combined.append(
                {
                    "id": str(obj.id),
                    "title": obj.get_localized_title(),
                    "description": obj.get_localized_description()[:150],
                    "resource_type": rtype,
                    "author": obj.author.get_full_name() if obj.author else "",
                    "author_email": obj.author.email if obj.author else "",
                    "date": obj.creation_date.strftime("%b %d, %Y")
                    if obj.creation_date
                    else "",
                    "url": reverse(
                        "resources:resource-detail",
                        kwargs={"type": rtype, "pk": obj.id},
                    ),
                }
            )

    if sort_by == "oldest":
        combined.sort(key=lambda x: x["date"])
    elif sort_by == "popular":
        pass  # already sorted by queryset order
    else:
        combined.sort(key=lambda x: x["date"], reverse=True)

    return JsonResponse({"resources": combined, "count": len(combined)})


class ToolListView(LoginAndVerifiedRequiredMixin, ListView):
    model = NLPTool
    template_name = "resources/tool_list.html"
    context_object_name = "tools"
    paginate_by = 10

    def get_queryset(self):
        # STRICT: Only show APPROVED tools in the public section
        # Pending tools are only visible in the admin panel
        queryset = exclude_hidden_users(
            NLPTool.objects.filter(approval_status="approved"),
            self.request.user,
            ("author",),
        )

        # Filter by tool type/category
        tool_type = self.request.GET.get("type", "").strip()
        if tool_type:
            queryset = queryset.filter(tool_type=tool_type)

        # Search functionality
        search_query = self.request.GET.get("q", "").strip()
        if search_query:
            queryset = queryset.filter(
                Q(title__icontains=search_query)
                | Q(title_ar__icontains=search_query)
                | Q(title_en__icontains=search_query)
                | Q(description__icontains=search_query)
                | Q(tool_type__icontains=search_query)
                | Q(keywords__icontains=search_query)
                | Q(author__first_name__icontains=search_query)
                | Q(author__last_name__icontains=search_query)
                | Q(supported_languages__icontains=search_query)
            ).distinct()

        return queryset.order_by("-creation_date")

    def get_template_names(self):
        """Return partial template for AJAX requests"""
        if self.request.headers.get("x-requested-with") == "XMLHttpRequest":
            return ["resources/_tool_grid_cards.html"]
        return [self.template_name]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        search_query = self.request.GET.get("q", "")
        tool_type = self.request.GET.get("type", "")

        # Always use filtered queryset for count (respects approval_status)
        context["total_count"] = self.get_queryset().count()

        if search_query:
            context["search_query"] = search_query
            context["is_search"] = True
        else:
            context["is_search"] = False

        # Tool type choices for filter chips
        context["tool_type_choices"] = NLPTool.ToolType.choices

        context["current_type"] = tool_type
        context["page"] = "tools"
        return context


class CourseListView(LoginAndVerifiedRequiredMixin, ListView):
    model = Course
    template_name = "resources/course_list.html"
    context_object_name = "courses"
    paginate_by = 10

    def get_template_names(self):
        # Return partial template for AJAX requests (live search)
        if self.request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return ["resources/partials/course_grid.html"]
        return [self.template_name]

    def get_queryset(self):
        # STRICT: Only show APPROVED courses in the public section
        # Pending courses are only visible in the admin panel
        queryset = exclude_hidden_users(
            Course.objects.filter(approval_status="approved"),
            self.request.user,
            ("author", "teacher"),
        )

        # Search filter
        search_query = self.request.GET.get("q", "").strip()
        if search_query:
            queryset = queryset.filter(
                Q(title__icontains=search_query)
                | Q(title_ar__icontains=search_query)
                | Q(title_en__icontains=search_query)
                | Q(description__icontains=search_query)
                | Q(keywords__icontains=search_query)
                | Q(author__first_name__icontains=search_query)
                | Q(author__last_name__icontains=search_query)
                | Q(field__icontains=search_query)
                | Q(academic_level__icontains=search_query)
                | Q(institution__name__icontains=search_query)
            ).distinct()

        # Level filter (Bachelor, Master, Doctorate)
        level = self.request.GET.get("level", "").strip().lower()
        if level in ["bachelor", "master", "doctorate"]:
            queryset = queryset.filter(academic_level=level)

        # Sort filter
        sort = self.request.GET.get("sort", "newest")
        if sort == "oldest":
            queryset = queryset.order_by("creation_date")
        elif sort == "alphabetical":
            queryset = queryset.order_by("title")
        elif sort == "popular":
            queryset = queryset.order_by("-views_count")
        else:  # newest (default)
            queryset = queryset.order_by("-creation_date")

        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        search_query = self.request.GET.get("q", "")
        level = self.request.GET.get("level", "")
        sort = self.request.GET.get("sort", "newest")

        # Always use filtered queryset for count (respects approval_status)
        context["total_count"] = self.get_queryset().count()

        if search_query:
            context["search_query"] = search_query
            context["is_search"] = True
        else:
            context["is_search"] = False

        context["current_level"] = level.lower() if level else ""
        context["current_sort"] = sort
        context["page"] = "course"
        return context


class ArticleListView(LoginAndVerifiedRequiredMixin, ListView):
    model = Article
    template_name = "resources/article_list.html"
    context_object_name = "articles"
    paginate_by = 10

    def get_queryset(self):
        # STRICT: Public articles list only shows approved content.
        if self.request.user.is_staff:
            return Article.objects.filter(approval_status="approved")
        return exclude_hidden_users(
            Article.objects.filter(approval_status="approved"),
            self.request.user,
            ("author",),
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["total_count"] = self.get_queryset().count()
        return context


class ThesisListView(LoginAndVerifiedRequiredMixin, ListView):
    model = Thesis
    template_name = "resources/thesis_list.html"
    context_object_name = "theses"
    paginate_by = 10

    def get_queryset(self):
        # STRICT: Public theses list only shows approved content.
        if self.request.user.is_staff:
            return Thesis.objects.filter(approval_status="approved")
        return exclude_hidden_users(
            Thesis.objects.filter(approval_status="approved"),
            self.request.user,
            ("author",),
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["total_count"] = self.get_queryset().count()
        return context


class MemoirListView(LoginAndVerifiedRequiredMixin, ListView):
    model = Memoir
    template_name = "resources/memoir_list.html"
    context_object_name = "memoirs"
    paginate_by = 10

    def get_queryset(self):
        # STRICT: Public memoirs list only shows approved content.
        if self.request.user.is_staff:
            return Memoir.objects.filter(approval_status="approved")
        return exclude_hidden_users(
            Memoir.objects.filter(approval_status="approved"),
            self.request.user,
            ("author",),
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["total_count"] = self.get_queryset().count()
        return context


class CorpusListView(LoginAndVerifiedRequiredMixin, ListView):
    model = Corpus
    template_name = "resources/corpus_list.html"
    context_object_name = "corpora"
    paginate_by = 10

    def get_template_names(self):
        """Return partial template for AJAX requests."""
        if self.request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return ["resources/_corpus_cards.html"]
        return [self.template_name]

    def get_queryset(self):
        # Staff can inspect all corpora; non-staff can only see approved corpora.
        if self.request.user.is_staff:
            queryset = Corpus.objects.all()
        else:
            queryset = Corpus.objects.filter(approval_status="approved")

        # Search query
        search_query = self.request.GET.get("q", "").strip()
        if search_query:
            queryset = queryset.filter(
                Q(title__icontains=search_query)
                | Q(title_ar__icontains=search_query)
                | Q(title_en__icontains=search_query)
                | Q(description__icontains=search_query)
                | Q(keywords__icontains=search_query)
                | Q(author__first_name__icontains=search_query)
                | Q(author__last_name__icontains=search_query)
                | Q(field__icontains=search_query)
            ).distinct()

        # Filter by fields (categories) - supports multiple values
        fields = self.request.GET.getlist("field")
        if fields:
            queryset = queryset.filter(field__in=fields)

        # Filter by languages - supports multiple values
        languages = self.request.GET.getlist("language")
        if languages:
            queryset = queryset.filter(language__in=languages)

        return queryset.order_by("-creation_date")

    def render_to_response(self, context, **response_kwargs):
        # Return partial HTML for AJAX requests
        if self.request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return render(self.request, "resources/_corpus_cards.html", context)
        return super().render_to_response(context, **response_kwargs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        search_query = self.request.GET.get("q", "")

        # Always use filtered queryset for count (respects approval_status)
        context["total_count"] = self.get_queryset().count()

        if search_query:
            context["search_query"] = search_query
            context["is_search"] = True
        else:
            context["is_search"] = False

        # Provide field choices for the filter sidebar
        from .models import FieldChoices

        context["field_choices"] = FieldChoices.choices

        # Track active filters
        context["active_fields"] = self.request.GET.getlist("field")
        context["active_languages"] = self.request.GET.getlist("language")

        # Check if any filters are active
        context["has_active_filters"] = bool(
            context["active_fields"] or context["active_languages"] or search_query
        )

        context["page"] = "corpus"
        return context


class ResourceDetailView(LoginAndVerifiedRequiredMixin, DetailView):
    template_name = "resources/resource_detail.html"
    context_object_name = "object"

    TYPE_MODELS: dict[str, type[models.Model]] = {
        "tool": NLPTool,
        "nlp_tool": NLPTool,
        "course": Course,
        "article": Article,
        "thesis": Thesis,
        "memoir": Memoir,
        "corpus": Corpus,
    }

    MODEL_VIEW_NAMES: dict[str, str] = {
        "nlptool": "tool",
        "course": "course",
        "article": "article",
        "thesis": "thesis",
        "memoir": "memoir",
        "corpus": "corpus",
    }

    URL_NAMES: dict[str, str] = {
        "tool": "tool_list",
        "course": "course_list",
        "article": "article_list",
        "thesis": "thesis_list",
        "memoir": "memoir_list",
        "corpus": "corpus_list",
        "document": "list",
    }

    def get_object(self):
        resource_type = self.kwargs.get("type")
        pk = self.kwargs.get("pk")
        is_staff_user = bool(
            self.request.user.is_authenticated
            and (self.request.user.is_staff or self.request.user.is_superuser)
        )

        model = self.TYPE_MODELS.get(resource_type)
        if not model:
            raise Http404("Type de ressource invalide")

        if resource_type in ["article", "thesis", "memoir"]:
            try:
                obj = get_object_or_404(model, pk=pk)
            except Http404 as err:
                document = get_object_or_404(Document, pk=pk)
                if resource_type == "article" and hasattr(document, "article"):
                    obj = document.article
                elif resource_type == "thesis" and hasattr(document, "thesis"):
                    obj = document.thesis
                elif resource_type == "memoir" and hasattr(document, "memoir"):
                    obj = document.memoir
                elif document.document_type == resource_type and is_staff_user:
                    # Some legacy documents can miss the related subtype row.
                    # Let staff/admin access the base Document to avoid blocking admin actions.
                    obj = document
                else:
                    raise Http404(
                        f"No {resource_type.capitalize()} matches the given query."
                    ) from err
        else:
            obj = get_object_or_404(model, pk=pk)

        # Staff/admin can access all statuses from admin/publication screens.
        # Non-staff users can only view approved resources.
        if not is_staff_user:
            if hasattr(obj, "approval_status"):
                approval_status = getattr(obj, "approval_status", None)
            else:
                related_document = getattr(obj, "document", None)
                approval_status = getattr(related_document, "approval_status", None)

            if approval_status is not None and approval_status != "approved":
                raise Http404("This resource is pending approval.")

        hidden_ids = blocked_user_ids_for(self.request.user)
        if hidden_ids:
            owner_id = getattr(obj, "author_id", None) or getattr(
                obj, "teacher_id", None
            )
            if owner_id in hidden_ids:
                raise Http404("Resource not found.")

        return obj

    def get(self, request, *args, **kwargs):
        self.object = self.get_object()

        resource_type = self.kwargs.get("type")

        if resource_type in ["article", "thesis", "memoir"]:
            if hasattr(self.object, "document") and self.object.document:
                self.object.document.increment_views()
            elif isinstance(self.object, Document):
                self.object.increment_views()
            else:
                logger.warning(f"Object {self.object.pk} has no associated document")
        elif hasattr(self.object, "increment_views"):
            self.object.increment_views()
        else:
            logger.warning(f"Object {self.object.pk} has no increment_views method")

        context = self.get_context_data(object=self.object)
        return self.render_to_response(context)

    def get_template_names(self):
        # In admin review mode, force the shared detail template so all
        # resource types (tools/corpus/articles/...) use the same approval UI.
        if (
            self.request.user.is_authenticated
            and (self.request.user.is_staff or self.request.user.is_superuser)
            and self.request.GET.get("admin_review") == "1"
        ):
            return [self.template_name]

        # If a staff user is viewing a document subtype with a missing related
        # subtype row, ResourceDetailView falls back to base Document.
        # Render with shared template to avoid subtype-template assumptions.
        if (
            self.request.user.is_authenticated
            and (self.request.user.is_staff or self.request.user.is_superuser)
            and self.kwargs.get("type") in ["article", "thesis", "memoir"]
            and isinstance(getattr(self, "object", None), Document)
        ):
            return [self.template_name]

        return [f"resources/{self.kwargs['type']}_detail.html", self.template_name]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        model_name = str(self.object._meta.model_name)
        resource_type = self.MODEL_VIEW_NAMES.get(model_name, model_name) or model_name
        context["resource_type"] = resource_type
        context["list_url_name"] = self.URL_NAMES.get(resource_type, "list")
        context["page"] = "resources"

        document_obj = getattr(self.object, "document", None)
        if document_obj is not None:
            context["specific_object"] = self.object
            context["object"] = document_obj

        if resource_type in ["article", "thesis", "memoir", "course"]:
            field = getattr(self.object, "field", None)
            if field is None and document_obj is not None:
                field = getattr(document_obj, "field", None)

            if field:
                context["related_corpora"] = Corpus.objects.filter(
                    approval_status="approved", field__icontains=field
                )[:3]
            else:
                context["related_corpora"] = Corpus.objects.filter(
                    approval_status="approved"
                )[:3]

        return context


class ResourceUpdateView(LoginRequiredMixin, UserPassesTestMixin, FormView):
    form_class = ResourceForm
    template_name = "resources/resource_update_form.html"

    # Updated TYPE_MODELS to treat article, thesis, memoir as top-level types
    TYPE_MODELS: dict[str, type[ResourceBase]] = {
        "tool": NLPTool,
        "nlp_tool": NLPTool,
        "course": Course,
        "corpus": Corpus,
        "article": Document,  # Changed: Article is accessed via Document
        "thesis": Document,  # Changed: Thesis is accessed via Document
        "memoir": Document,  # Changed: Memoir is accessed via Document
    }

    def get_object(self) -> ResourceBase:
        resource_type = self.kwargs["type"]
        pk = self.kwargs["pk"]
        is_staff_user = bool(
            self.request.user.is_authenticated
            and (self.request.user.is_staff or self.request.user.is_superuser)
        )

        # Get the Document for article, thesis, memoir
        if resource_type in ["article", "thesis", "memoir"]:
            document = get_object_or_404(Document, pk=pk)
            # For staff/admin, always allow editing by document type even if the
            # related subtype row is missing (legacy/inconsistent rows).
            if is_staff_user and document.document_type == resource_type:
                return document

            # Verify the document has the correct subtype; if not, try to find the actual subtype
            if not hasattr(document, resource_type):
                # Fallback: check other subtypes
                for subtype in ["article", "thesis", "memoir"]:
                    if hasattr(document, subtype):
                        return document

                if document.document_type == resource_type:
                    return document

                raise Http404(
                    f"{resource_type.capitalize()} not found for document ID {pk}"
                )
            return document
        else:
            model = self.TYPE_MODELS.get(resource_type)
            if not model:
                raise Http404("Invalid resource type")
            return get_object_or_404(model, pk=pk)

    def get_initial(self):
        resource = cast(ResourceBase, self.get_object())
        resource_type = self.kwargs["type"]
        initial = {}

        # Common fields
        initial.update(
            {
                "title_en": resource.title_en or resource.title,
                "title_ar": resource.title_ar or "",
                "description_en": resource.description_en or resource.description,
                "description_ar": resource.description_ar or "",
                "keywords": resource.keywords,
                "access_link": resource.access_link or "",
                "language": resource.language,
            }
        )

        # Type-specific fields
        if resource_type == "course" and isinstance(resource, Course):
            initial.update(
                {
                    "course_field": resource.field,
                    "academic_level": resource.academic_level,
                    "course_institution": resource.institution.id
                    if resource.institution
                    else None,
                    "academic_year": resource.academic_year,
                    "prerequisites": resource.prerequisites,
                    "syllabus": resource.syllabus,
                    "resource_type": "course",
                }
            )
        elif resource_type in ["nlp_tool", "tool"] and isinstance(resource, NLPTool):
            initial.update(
                {
                    "tool_type": resource.tool_type,
                    "tool_version": resource.version,
                    "documentation": resource.documentation_link or "",
                    "supported_languages": resource.get_supported_languages_list()
                    if hasattr(resource, "get_supported_languages_list")
                    else [],
                    "resource_type": "nlp_tool",
                }
            )
        elif resource_type == "corpus" and isinstance(resource, Corpus):
            initial.update({"corpus_field": resource.field, "resource_type": "corpus"})
        elif resource_type == "article":
            article = getattr(resource, "article", None)
            if article:
                initial.update(
                    {
                        "document_format": getattr(resource, "file_format", ""),
                        "journal": article.journal,
                        "publication_date": article.publication_date,
                        "doi": article.doi or "",
                        "resource_type": "article",
                    }
                )
        elif resource_type == "thesis":
            thesis = getattr(resource, "thesis", None)
            if thesis:
                initial.update(
                    {
                        "document_format": getattr(resource, "file_format", ""),
                        "supervisor": thesis.supervisor,
                        "thesis_institution": thesis.institution.id
                        if thesis.institution
                        else None,
                        "defense_year": thesis.defense_year,
                        "resource_type": "thesis",
                    }
                )
        elif resource_type == "memoir":
            memoir = getattr(resource, "memoir", None)
            if memoir:
                initial.update(
                    {
                        "document_format": getattr(resource, "file_format", ""),
                        "memoir_level": memoir.academic_level,
                        "memoir_institution": memoir.institution.id
                        if memoir.institution
                        else None,
                        "memoir_defense_year": memoir.defense_year,
                        "resource_type": "memoir",
                    }
                )

        return initial

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["user"] = self.request.user
        kwargs["is_update"] = True
        kwargs["instance"] = self.get_object()
        return kwargs

    def form_valid(self, form):
        resource = cast(ResourceBase, self.get_object())
        resource_type = self.kwargs["type"]
        review_mode = (
            self.request.GET.get("review") == "1" and self.request.user.is_staff
        )
        edit_only = (
            self.request.GET.get("edit_only") == "1" and self.request.user.is_staff
        )

        # In review view mode, form must be read-only. Only approve/reject actions are allowed.
        if review_mode and not edit_only:
            messages.warning(
                self.request,
                _(
                    "This form is read-only in review mode. Use Edit mode to modify fields."
                ),
            )
            return redirect(self.request.get_full_path())

        # Use form.save() which correctly handles all fields and type-specific logic
        try:
            resource = form.save(instance=resource)
        except Exception as e:
            logger.warning("Error during resource save (may be ES indexing): %s", e)
            resource.refresh_from_db()

        # Handle "Approve & Publish" button
        if self.request.POST.get("approve_and_publish") and self.request.user.is_staff:
            # Validate bilingual fields before approving
            missing = []
            if not resource.title_ar:
                missing.append(_("Title (Arabic)"))
            if not resource.title_en:
                missing.append(_("Title (English)"))
            if not resource.description_ar:
                missing.append(_("Description (Arabic)"))
            if not resource.description_en:
                missing.append(_("Description (English)"))

            if missing:
                messages.error(
                    self.request,
                    _("Cannot approve: Missing %(fields)s")
                    % {"fields": ", ".join(str(f) for f in missing)},
                )
                return redirect(self.request.get_full_path())

            resource.approval_status = "approved"
            try:
                resource.save(update_fields=["approval_status"])
            except Exception as e:
                logger.warning(
                    "ES indexing error during resource approval (saved OK): %s", e
                )

            # Notify author
            from notifications.services import NotificationService

            author = resource.author
            if author:
                NotificationService.create_notification(
                    recipient=author,
                    notification_type="POST_APPROVED",
                    title=_("Your submission has been approved"),
                    message=_(
                        "Your submission '%(title)s' has been approved and is now visible to the public."
                    ),
                    message_kwargs={"title": resource.title},
                )

            messages.success(
                self.request,
                _("'%(title)s' has been approved and published!")
                % {"title": resource.title},
            )
            # Redirect to admin page instead of detail page
            return redirect(self.get_admin_redirect_url(resource_type))

        messages.success(
            self.request,
            _("Resource '%(title)s' updated successfully!") % {"title": resource.title},
        )

        # In admin edit-only mode, go back to the review detail page to allow approve/reject.
        if (
            edit_only
            and self.request.GET.get("review_model")
            and self.request.GET.get("review_pk")
        ):
            return redirect(self.request.get_full_path())

        # In review mode, redirect back to admin
        if self.request.GET.get("review") == "1" and self.request.user.is_staff:
            messages.success(self.request, _("Draft saved successfully."))
            return redirect(self.get_admin_redirect_url(resource_type))

        return super().form_valid(form)

    def get_admin_redirect_url(self, resource_type):
        """Get the admin page URL for the resource type."""
        admin_urls = {
            "course": "pages:admin_courses",
            "tool": "pages:admin_tools",
            "nlp_tool": "pages:admin_tools",
            "corpus": "pages:admin_corpora",
            "article": "pages:admin_publications",
            "thesis": "pages:admin_publications",
            "memoir": "pages:admin_publications",
        }
        url_name = admin_urls.get(resource_type, "pages:admin_dashboard")
        return f"{reverse(url_name)}?tab=pending"

    def get_success_url(self):
        resource_type = self.kwargs["type"]
        pk = self.kwargs["pk"]
        return reverse(
            "resources:resource-detail", kwargs={"type": resource_type, "pk": pk}
        )

    def test_func(self):
        if self.request.user.is_staff or self.request.user.is_superuser:
            return True
        resource = self.get_object()
        return resource.author == self.request.user

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["page"] = "resources"
        # Check if in admin review mode
        context["review_mode"] = self.request.GET.get("review") == "1"
        context["edit_only"] = self.request.GET.get("edit_only") == "1"
        resource = self.get_object()
        context["is_pending"] = getattr(resource, "approval_status", None) == "pending"
        context["resource"] = resource
        return context


class ResourceDeleteView(
    LoginAndVerifiedRequiredMixin, UserPassesTestMixin, DeleteView
):
    template_name = "resources/resource_confirm_delete.html"
    success_url = reverse_lazy("resources:list")

    TYPE_MODELS = {
        "tool": NLPTool,
        "nlp_tool": NLPTool,
        "course": Course,
        "corpus": Corpus,
        "article": Document,
        "thesis": Document,
        "memoir": Document,
    }

    def get_object(self):
        model = self.TYPE_MODELS.get(self.kwargs["type"])
        if not model:
            raise Http404("Invalid resource type")
        return get_object_or_404(model, pk=self.kwargs["pk"])

    def delete(self, request, *args, **kwargs):
        resource = self.get_object()
        resource_title = resource.title

        # For Document types, the related Article/Thesis/Memoir will be deleted automatically
        # due to OneToOneField cascade
        response = super().delete(request, *args, **kwargs)

        messages.success(
            self.request, f"Resource '{resource_title}' deleted successfully!"
        )
        return response

    def test_func(self):
        resource = self.get_object()
        if self.request.user.is_staff or self.request.user.is_superuser:
            return True
        return resource.author == self.request.user

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["page"] = "resources"
        return context


class ResourceCreateView(LoginAndVerifiedRequiredMixin, FormView):
    template_name = "resources/resource_form.html"
    form_class = ResourceForm
    success_url = reverse_lazy("resources:list")

    def post(self, request, *args, **kwargs):
        import logging

        logger = logging.getLogger(__name__)
        logger.info(
            f"[RESOURCE_CREATE] POST request received from user: {request.user.email}"
        )
        logger.info(f"[RESOURCE_CREATE] POST data keys: {list(request.POST.keys())}")
        logger.info(f"[RESOURCE_CREATE] FILES data keys: {list(request.FILES.keys())}")

        # Log critical fields
        logger.info(
            f"[RESOURCE_CREATE] resource_type: {request.POST.get('resource_type')}"
        )
        logger.info(f"[RESOURCE_CREATE] title_en: {request.POST.get('title_en')}")
        logger.info(f"[RESOURCE_CREATE] title_ar: {request.POST.get('title_ar')}")

        return super().post(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["user"] = self.request.user
        return kwargs

    def form_valid(self, form):
        import logging

        logger = logging.getLogger(__name__)
        logger.info(
            f"[RESOURCE_CREATE] Form is valid, attempting to save for user: {self.request.user.email}"
        )

        try:
            resource = form.save()
            logger.info(
                f"[RESOURCE_CREATE] ✓ Resource saved successfully (ID: {resource.id})"
            )
            if (
                self.request.user.is_staff
                and getattr(resource, "approval_status", None) != "approved"
            ):
                resource.approval_status = "approved"
                resource.save(update_fields=["approval_status"])
                messages.success(
                    self.request,
                    _("Your submission '%(title)s' has been created and published.")
                    % {"title": resource.title},
                )
            else:
                messages.info(
                    self.request,
                    _(
                        "Your submission '%(title)s' has been received and is pending admin review. It will be visible to the public once approved."
                    )
                    % {"title": resource.title},
                )
            return super().form_valid(form)
        except Exception as e:
            logger.error(
                f"[RESOURCE_CREATE] ✗ Error saving resource: {str(e)}", exc_info=True
            )
            from django.conf import settings

            if settings.DEBUG:
                messages.error(self.request, f"Error: {str(e)}")
            else:
                messages.error(
                    self.request,
                    _(
                        "An error occurred while creating the resource. Please try again."
                    ),
                )
            return self.form_invalid(form)

    def form_invalid(self, form):
        import logging

        logger = logging.getLogger(__name__)
        logger.warning(
            f"[RESOURCE_CREATE] ✗ Form validation failed for user: {self.request.user.email}"
        )
        logger.warning(f"[RESOURCE_CREATE] Form errors: {form.errors.as_json()}")
        logger.warning(f"[RESOURCE_CREATE] Non-field errors: {form.non_field_errors()}")

        # Log which fields have errors
        for field, errors in form.errors.items():
            logger.warning(f"[RESOURCE_CREATE] Field '{field}' errors: {errors}")

        messages.error(
            self.request, _("Please correct the errors in the form and try again.")
        )
        return super().form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["page"] = "resources"
        return context


class CourseCreateView(LoginAndVerifiedRequiredMixin, FormView):
    template_name = "resources/course_create_form.html"
    form_class = ResourceForm
    success_url = reverse_lazy("resources:course_list")

    def post(self, request, *args, **kwargs):
        import logging

        logger = logging.getLogger(__name__)
        logger.info(
            f"[COURSE_CREATE] POST request received from user: {request.user.email}"
        )
        logger.info(f"[COURSE_CREATE] POST data keys: {list(request.POST.keys())}")
        logger.info(f"[COURSE_CREATE] FILES data keys: {list(request.FILES.keys())}")

        # Log critical fields
        logger.info(
            f"[COURSE_CREATE] resource_type: {request.POST.get('resource_type')}"
        )
        logger.info(f"[COURSE_CREATE] title_en: {request.POST.get('title_en')}")
        logger.info(f"[COURSE_CREATE] title_ar: {request.POST.get('title_ar')}")

        return super().post(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["user"] = self.request.user
        return kwargs

    def get_initial(self):
        initial = super().get_initial()
        initial["resource_type"] = "course"
        return initial

    def form_valid(self, form):
        import logging

        logger = logging.getLogger(__name__)
        logger.info(
            f"[COURSE_CREATE] Form is valid, attempting to save for user: {self.request.user.email}"
        )

        try:
            resource = form.save()
            if (
                self.request.user.is_staff
                and getattr(resource, "approval_status", None) != "approved"
            ):
                resource.approval_status = "approved"
                resource.save(update_fields=["approval_status"])
                logger.info(
                    f"[COURSE_CREATE] Auto-approved staff course (ID: {resource.id})"
                )
                messages.success(
                    self.request,
                    _("Your course '%(title)s' has been created and published.")
                    % {"title": resource.title},
                )
            else:
                logger.info(
                    f"[COURSE_CREATE] Course saved successfully (ID: {resource.id})"
                )
                messages.info(
                    self.request,
                    _(
                        "Your course '%(title)s' has been submitted and is pending admin review."
                    )
                    % {"title": resource.title},
                )
            return super().form_valid(form)
        except Exception as e:
            logger.error(
                f"[COURSE_CREATE] Error saving course: {str(e)}", exc_info=True
            )
            messages.error(
                self.request,
                _("An error occurred while creating the course. Please try again."),
            )
            return self.form_invalid(form)

    def form_invalid(self, form):
        import logging

        logger = logging.getLogger(__name__)
        logger.warning(
            f"[COURSE_CREATE] ✗ Form validation failed for user: {self.request.user.email}"
        )
        logger.warning(f"[COURSE_CREATE] Form errors: {form.errors.as_json()}")
        logger.warning(f"[COURSE_CREATE] Non-field errors: {form.non_field_errors()}")

        # Log which fields have errors
        for field, errors in form.errors.items():
            logger.warning(f"[COURSE_CREATE] Field '{field}' errors: {errors}")

        messages.error(
            self.request, _("Please correct the errors in the form and try again.")
        )
        return super().form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["page"] = "resources"
        return context


class CorpusCreateView(LoginAndVerifiedRequiredMixin, FormView):
    template_name = "resources/corpus_create_form.html"
    form_class = ResourceForm
    success_url = reverse_lazy("resources:corpus_list")

    def post(self, request, *args, **kwargs):
        import logging

        logger = logging.getLogger(__name__)
        logger.info(
            f"[CORPUS_CREATE] POST request received from user: {request.user.email}"
        )
        logger.info(f"[CORPUS_CREATE] POST data keys: {list(request.POST.keys())}")
        logger.info(f"[CORPUS_CREATE] FILES data keys: {list(request.FILES.keys())}")

        # Log critical fields
        logger.info(
            f"[CORPUS_CREATE] resource_type: {request.POST.get('resource_type')}"
        )
        logger.info(f"[CORPUS_CREATE] title_en: {request.POST.get('title_en')}")
        logger.info(f"[CORPUS_CREATE] title_ar: {request.POST.get('title_ar')}")

        return super().post(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["user"] = self.request.user
        return kwargs

    def get_initial(self):
        initial = super().get_initial()
        initial["resource_type"] = "corpus"
        return initial

    def form_valid(self, form):
        import logging

        logger = logging.getLogger(__name__)
        logger.info(
            f"[CORPUS_CREATE] Form is valid, attempting to save for user: {self.request.user.email}"
        )

        try:
            resource = form.save()
            logger.info(
                f"[CORPUS_CREATE] ✓ Corpus saved successfully (ID: {resource.id})"
            )
            success_message = _(
                "Your corpus '%(title)s' has been created successfully and is pending admin approval."
            ) % {"title": resource.get_localized_title()}
            if self.request.headers.get("X-Requested-With") == "XMLHttpRequest":
                return JsonResponse(
                    {
                        "success": True,
                        "message": str(success_message),
                        "redirect_url": str(self.get_success_url()),
                    }
                )
            messages.success(self.request, success_message)
            return super().form_valid(form)
        except Exception as e:
            logger.error(
                f"[CORPUS_CREATE] ✗ Error saving corpus: {str(e)}", exc_info=True
            )
            messages.error(
                self.request,
                _("An error occurred while creating the corpus. Please try again."),
            )
            return self.form_invalid(form)

    def form_invalid(self, form):
        import logging

        logger = logging.getLogger(__name__)
        logger.warning(
            f"[CORPUS_CREATE] ✗ Form validation failed for user: {self.request.user.email}"
        )
        logger.warning(f"[CORPUS_CREATE] Form errors: {form.errors.as_json()}")
        logger.warning(f"[CORPUS_CREATE] Non-field errors: {form.non_field_errors()}")

        # Log which fields have errors
        for field, errors in form.errors.items():
            logger.warning(f"[CORPUS_CREATE] Field '{field}' errors: {errors}")

        messages.error(
            self.request, _("Please correct the errors in the form and try again.")
        )
        if self.request.headers.get("X-Requested-With") == "XMLHttpRequest":
            return JsonResponse(
                {
                    "success": False,
                    "errors": form.errors,
                    "non_field_errors": form.non_field_errors(),
                },
                status=400,
            )
        return super().form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["page"] = "resources"
        return context


class ToolCreateView(LoginAndVerifiedRequiredMixin, FormView):
    template_name = "resources/tool_create_form.html"
    form_class = ResourceForm
    success_url = reverse_lazy("resources:tool_list")

    def post(self, request, *args, **kwargs):
        import logging

        logger = logging.getLogger(__name__)
        logger.info(
            f"[TOOL_CREATE] POST request received from user: {request.user.email}"
        )
        logger.info(f"[TOOL_CREATE] POST data keys: {list(request.POST.keys())}")
        logger.info(f"[TOOL_CREATE] FILES data keys: {list(request.FILES.keys())}")

        # Log critical fields
        logger.info(f"[TOOL_CREATE] resource_type: {request.POST.get('resource_type')}")
        logger.info(f"[TOOL_CREATE] title_en: {request.POST.get('title_en')}")
        logger.info(f"[TOOL_CREATE] title_ar: {request.POST.get('title_ar')}")

        return super().post(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs["user"] = self.request.user
        return kwargs

    def get_initial(self):
        initial = super().get_initial()
        initial["resource_type"] = "nlp_tool"
        return initial

    def form_valid(self, form):
        import logging

        logger = logging.getLogger(__name__)
        logger.info(
            f"[TOOL_CREATE] Form is valid, attempting to save for user: {self.request.user.email}"
        )

        try:
            resource = form.save()
            logger.info(f"[TOOL_CREATE] ✓ Tool saved successfully (ID: {resource.id})")
            messages.info(
                self.request,
                _(
                    "Your tool '%(title)s' has been submitted and is pending admin review."
                )
                % {"title": resource.title},
            )
            return super().form_valid(form)
        except Exception as e:
            logger.error(f"[TOOL_CREATE] ✗ Error saving tool: {str(e)}", exc_info=True)
            messages.error(
                self.request,
                _("An error occurred while creating the tool. Please try again."),
            )
            return self.form_invalid(form)

    def form_invalid(self, form):
        import logging

        logger = logging.getLogger(__name__)
        logger.warning(
            f"[TOOL_CREATE] ✗ Form validation failed for user: {self.request.user.email}"
        )
        logger.warning(f"[TOOL_CREATE] Form errors: {form.errors.as_json()}")
        logger.warning(f"[TOOL_CREATE] Non-field errors: {form.non_field_errors()}")

        # Log which fields have errors
        for field, errors in form.errors.items():
            logger.warning(f"[TOOL_CREATE] Field '{field}' errors: {errors}")

        messages.error(
            self.request, _("Please correct the errors in the form and try again.")
        )
        return super().form_invalid(form)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context["page"] = "resources"
        return context


# =============================================================================
# Convert to Text — PDF / DOCX / TXT text extraction (with OCR fallback)
# =============================================================================


def _extract_text_from_pdf(file_path: str) -> str:
    """
    Extract text from PDF using PyMuPDF.
    Falls back to Tesseract OCR for image-based pages.
    """
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    pages_text: list[str] = []

    for page_num in range(len(doc)):
        page = doc[page_num]
        text = page.get_text("text")

        # If a page yields very little text, it's likely an image-based page → OCR
        if len(text.strip()) < 30:
            text = _ocr_page(page) or text

        if text.strip():
            pages_text.append(f"--- Page {page_num + 1} ---\n{text.strip()}")

    doc.close()
    return "\n\n".join(pages_text)


def _ocr_page(page) -> str | None:
    """Run Tesseract OCR on a PDF page rendered as an image."""
    try:
        import io

        import pytesseract
        from PIL import Image

        # Render page at 300 DPI for good OCR quality
        pix = page.get_pixmap(dpi=300)
        img = Image.open(io.BytesIO(pix.tobytes("png")))

        # Support Arabic + English + French
        text = pytesseract.image_to_string(img, lang="ara+eng+fra")
        return text.strip() if text.strip() else None
    except Exception as e:
        logger.warning(f"OCR failed for page: {e}")
        return None


def _extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx."""
    from docx import Document as DocxDocument

    doc = DocxDocument(file_path)
    paragraphs: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                paragraphs.append(row_text)

    return "\n\n".join(paragraphs)


def _extract_text_from_txt(file_path: str) -> str:
    """Read plain text files with encoding detection."""
    for encoding in ("utf-8", "utf-8-sig", "cp1256", "iso-8859-6", "latin-1"):
        try:
            with open(file_path, encoding=encoding) as f:
                return f.read()
        except (UnicodeDecodeError, LookupError):
            continue
    # Last resort — binary read and ignore errors
    with open(file_path, encoding="utf-8", errors="ignore") as f:
        return f.read()


@login_required
@require_GET
def convert_to_text(request, pk):
    """
    API endpoint: extract raw text from a resource's uploaded file.
    Supports PDF (with OCR fallback), DOCX, and plain text files.

    Returns JSON: { success, text, filename, pages } or { success, error }
    """
    # Find the resource across all concrete models
    resource = None
    for model_cls in (Course, Corpus, Document, NLPTool, Article, Thesis, Memoir):
        try:
            resource = model_cls.objects.get(pk=pk)
            break
        except model_cls.DoesNotExist:
            continue

    if resource is None:
        return JsonResponse(
            {"success": False, "error": _("Resource not found.")}, status=404
        )

    if not resource.uploaded_file:
        return JsonResponse(
            {"success": False, "error": _("This resource has no uploaded file.")},
            status=400,
        )

    file_path = resource.uploaded_file.path
    if not os.path.isfile(file_path):
        return JsonResponse(
            {
                "success": False,
                "error": _("The file could not be found on the server."),
            },
            status=404,
        )

    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    try:
        if ext == ".pdf":
            text = _extract_text_from_pdf(file_path)
        elif ext in (".docx", ".doc"):
            text = _extract_text_from_docx(file_path)
        elif ext in (".txt", ".md", ".csv", ".json", ".xml", ".log"):
            text = _extract_text_from_txt(file_path)
        else:
            return JsonResponse(
                {
                    "success": False,
                    "error": _("Unsupported file format: %(ext)s") % {"ext": ext},
                },
                status=400,
            )

        if not text or not text.strip():
            return JsonResponse(
                {
                    "success": False,
                    "error": _(
                        "No text could be extracted from this document. It may be empty or contain only images without recognisable text."
                    ),
                },
                status=200,
            )

        return JsonResponse(
            {
                "success": True,
                "text": text,
                "filename": filename,
                "char_count": len(text),
                "word_count": len(text.split()),
            }
        )

    except Exception as e:
        logger.exception(f"Text extraction failed for resource {pk}")
        return JsonResponse(
            {
                "success": False,
                "error": _("An error occurred while processing the document: %(err)s")
                % {"err": str(e)},
            },
            status=500,
        )
