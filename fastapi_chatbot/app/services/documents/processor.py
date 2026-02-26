"""
Document processor — text extraction and intelligent chunking.

Supports PDF, TXT, DOCX and XLSX.  Chunks are sized to fit within
embedding model context windows while preserving sentence boundaries.
"""

import io
import logging
import re
from typing import List, Dict, Optional

import PyPDF2

from app.config import get_settings

logger = logging.getLogger(__name__)
settings = get_settings()


class DocumentProcessor:
    """Extract text from files and split into overlapping chunks."""

    def __init__(
        self,
        chunk_size: int = settings.CHUNK_SIZE,
        chunk_overlap: int = settings.CHUNK_OVERLAP,
        max_chunks: int = settings.MAX_CHUNKS_PER_DOC,
    ):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.max_chunks = max_chunks

    # ------------------------------------------------------------------
    # Text extraction
    # ------------------------------------------------------------------

    def extract_pdf_text(self, pdf_bytes: bytes) -> List[Dict]:
        """Extract text from PDF, returning list of {page, content}."""
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
        pages: List[Dict] = []
        for idx, page in enumerate(reader.pages, start=1):
            text = page.extract_text() or ""
            text = self._normalise(text)
            if text.strip():
                pages.append({"page": idx, "content": text})
        return pages

    def extract_docx_text(self, docx_bytes: bytes) -> List[Dict]:
        """Extract text from a DOCX file, returning list of {page, content}.

        Each paragraph becomes part of a single logical 'page' because
        DOCX files don't have physical page breaks we can reliably detect.
        We split on every ~3000 chars to create manageable chunks.
        """
        from docx import Document as DocxDocument

        doc = DocxDocument(io.BytesIO(docx_bytes))
        full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    full_text += "\n" + row_text

        full_text = self._normalise(full_text)
        if not full_text.strip():
            return []

        # Split into virtual pages (~3000 chars each)
        pages: List[Dict] = []
        chunk_size = 3000
        for i in range(0, len(full_text), chunk_size):
            segment = full_text[i : i + chunk_size].strip()
            if segment:
                pages.append({"page": len(pages) + 1, "content": segment})
        return pages

    def extract_xlsx_text(self, xlsx_bytes: bytes) -> List[Dict]:
        """Extract text from an XLSX file, one 'page' per sheet."""
        import openpyxl

        wb = openpyxl.load_workbook(
            io.BytesIO(xlsx_bytes), read_only=True, data_only=True
        )
        pages: List[Dict] = []
        for idx, sheet_name in enumerate(wb.sheetnames, start=1):
            ws = wb[sheet_name]
            rows_text: List[str] = []
            for row in ws.iter_rows(values_only=True):
                cells = [
                    str(c).strip() for c in row if c is not None and str(c).strip()
                ]
                if cells:
                    rows_text.append("\t".join(cells))
            if rows_text:
                content = f"Sheet: {sheet_name}\n" + "\n".join(rows_text)
                pages.append({"page": idx, "content": self._normalise(content)})
        wb.close()
        return pages

    def extract_text(self, raw_text: str) -> List[Dict]:
        """Wrap raw text string as a single page."""
        return [{"page": 1, "content": self._normalise(raw_text)}]

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def chunk_text(
        self,
        text: str,
        page_map: Optional[List[Dict]] = None,
    ) -> List[Dict]:
        """Split text into overlapping chunks.

        Returns list of {"content": str, "page": int | None}.
        """
        if page_map:
            return self._chunk_pages(page_map)
        return self._chunk_flat(text)

    def chunk_pdf(self, pdf_bytes: bytes) -> List[Dict]:
        pages = self.extract_pdf_text(pdf_bytes)
        return self._chunk_pages(pages)

    # ------------------------------------------------------------------
    # Internal
    # ------------------------------------------------------------------

    def _chunk_pages(self, pages: List[Dict]) -> List[Dict]:
        chunks: List[Dict] = []
        for page_info in pages:
            page_chunks = self._split(page_info["content"])
            for c in page_chunks:
                chunks.append({"content": c, "page": page_info["page"]})
                if len(chunks) >= self.max_chunks:
                    return chunks
        return chunks

    def _chunk_flat(self, text: str) -> List[Dict]:
        parts = self._split(text)
        return [{"content": p, "page": None} for p in parts[: self.max_chunks]]

    def _split(self, text: str) -> List[str]:
        """Sliding-window sentence-aware splitter.

        Tries to break on sentence boundaries ('.', '!', '?', Arabic period)
        then falls back to word boundaries.
        """
        text = text.strip()
        if not text:
            return []

        sentences = re.split(r"(?<=[.!?。؟])\s+", text)
        chunks: List[str] = []
        current: List[str] = []
        current_len = 0

        for sent in sentences:
            sent_len = len(sent.split())
            if current_len + sent_len > self.chunk_size and current:
                chunks.append(" ".join(current))
                # Overlap: keep tail sentences
                overlap_words = 0
                overlap_start = len(current)
                for i in range(len(current) - 1, -1, -1):
                    overlap_words += len(current[i].split())
                    if overlap_words >= self.chunk_overlap:
                        overlap_start = i
                        break
                current = current[overlap_start:]
                current_len = sum(len(s.split()) for s in current)

            current.append(sent)
            current_len += sent_len

        if current:
            chunks.append(" ".join(current))

        return chunks

    @staticmethod
    def clean_text(text: str) -> str:
        """Clean raw text before chunking.

        Phase 7: centralised cleaning step applied to all user-uploaded
        documents (both PDF-extracted and plain-text uploads).
        """
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", text)
        text = re.sub(r"[ \t]+", " ", text)
        text = re.sub(r"\n{3,}", "\n\n", text)
        return text.strip()

    @staticmethod
    def _normalise(text: str) -> str:
        """Basic whitespace normalisation."""
        return re.sub(r"\s+", " ", text).strip()


# Singleton
_processor: Optional[DocumentProcessor] = None


def get_document_processor() -> DocumentProcessor:
    global _processor
    if _processor is None:
        _processor = DocumentProcessor()
    return _processor
